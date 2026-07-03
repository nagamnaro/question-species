"""Update MVP Plan Section 13 with current implementation status."""

from docx import Document

PATH = r"C:\Users\oranm\OneDrive - Imperial College London\Question Species\Question Species MVP Plan.docx"


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(f"• {item}")


def main() -> None:
    doc = Document(PATH)

    doc.add_page_break()
    doc.add_heading("13. Implementation Status & Remaining Work", level=1)
    doc.add_paragraph(
        "Updated: July 2026. Reflects codebase after consistency fixes, "
        "profile personalisation, and social depth pass."
    )

    doc.add_heading("13.1 Build phases", level=2)
    for line in [
        "Phase 1 (Auth, feed, answer-before-reveal, basic results): Done",
        "Phase 2 (Species, friends, profile, submission): Done",
        "Phase 3 (AI clustering + insight summaries): Done (Groq + heuristic fallback)",
        "Phase 4 (Prediction, Mind Match, anti-echo, social layer): Done",
        "Phase 5 (Profile depth, moderation stub, content scale): Mostly done",
    ]:
        doc.add_paragraph(f"• {line}")

    doc.add_heading("13.2 Completed (aligned with spec)", level=2)
    add_bullets(
        doc,
        [
            "Magic-link auth, feed, species tabs, anti-echo interleaving (All + species tabs)",
            "Feed signals: response count, friend %, trending/popular/controversial/challenge badges",
            "Interactive upvotes, public replies/reactions with counts, private notes + inbox",
            "Answer-before-reveal, lock-in copy, optional crowd-prediction (prediction, opinion, discourse-tagged)",
            "Results: comparison (majority/minority/polarised), diverse views banner, echo prompt wired to insights",
            "Species scoring: puzzle, estimation median, prediction/crowd accuracy cards",
            "AI insight summary (Groq + heuristic), cached; insight backfill script",
            "Mind Match, agreement map, thought partners, Start discussion deep links",
            "Profile: bio, avatar upload, AI thinking tags (cached) with heuristic fallback",
            "Question submission with auto-tag/species suggest + quality scoring (pending vs published)",
            "Anti-echo fallback clusters when insights cache empty",
            "~258 seeded questions (after latest migration batch)",
        ],
    )

    doc.add_heading("13.3 Remaining — high priority", level=2)
    add_bullets(
        doc,
        [
            "Apply all Supabase migrations (see supabase/MIGRATIONS.md)",
            "Run npm run backfill-insights after deploy when responses exist",
            "Rich input modes: MCQ for opinion, sliders for estimation (input_config column deferred)",
            "Embedding-based clustering (currently JSON/heuristic clusters)",
            "Admin UI to approve pending submissions",
        ],
    )

    doc.add_heading("13.4 Remaining — polish", level=2)
    add_bullets(
        doc,
        [
            "Prediction ground truth vs real-world outcomes (explicitly later phase)",
            "Dedicated REST API layer (Supabase + server actions sufficient for MVP)",
            "Expand question bank toward 500",
            "Soft user profiling dimensions (partial via AI thinking tags)",
        ],
    )

    doc.save(PATH)
    print("Saved successfully")


if __name__ == "__main__":
    main()
